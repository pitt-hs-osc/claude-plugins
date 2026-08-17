#!/usr/bin/env python3
"""
upmc-release-to-web transform engine.

Two phases, so the writer always approves BEFORE anything is written:

  1) detect  - read the UPMC .docx, print a categorized change list, and write
               a machine-readable plan (<input>.plan.json). Nothing is changed.

  2) apply   - read that plan + the writer's approved categories, and write an
               edited COPY (YYYY-MM-DD-<slug>-web.docx). The original is never
               touched.

Rules live in references/upmc-to-pitt-rules.md and references/pitt-hs-style.md
(human-readable). The runtime copy of those rules is the RULES block below; keep
the two in sync.

Usage:
  python transform.py detect "input.docx" [--out plan.json]
  python transform.py apply  "input.docx" --plan plan.json \
        --approve STRIP_CONTACT,STRIP_EMBARGO,... [--slug my-slug] [--outdir DIR]
  python transform.py apply  "input.docx" --plan plan.json --approve ALL
"""

import argparse, json, os, re, sys, datetime, copy

try:
    import docx
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    sys.exit("python-docx is required: pip install python-docx --break-system-packages")

# --------------------------------------------------------------------------- #
# RULES (runtime mirror of references/*.md)                                    #
# --------------------------------------------------------------------------- #

# Paragraph-strip rules: (category, label, matcher)
def _starts(*prefixes):
    return lambda t: any(t.lower().startswith(p.lower()) for p in prefixes)

PARA_STRIP_RULES = [
    ("STRIP_CONTACT",      "Media contact lines",        _starts("Contact:", "Mobile:", "E-mail:", "Email:")),
    ("STRIP_EMBARGO",      "Embargo line",               lambda t: "embargoed" in t.lower()),
    ("STRIP_SUMMARY",      "'Summary:' label",           _starts("Summary:")),
    ("STRIP_HASHES",       "End-of-release marker",      lambda t: re.fullmatch(r"[#\s]+|-\s*30\s*-", t.strip()) is not None and "#" in t or t.strip() in ("-30-",)),
    ("STRIP_MEDIA_FOOTER", "UPMC media footer",          lambda t: re.sub(r"^https?://", "", t.strip().lower()).rstrip("/") == "www.upmc.com/media"),
]

# Block-strip rules: start matcher -> remove start para + following body paras
# until a stop condition. Used for the "About …" boilerplate and "Additional
# Resources" list.
def _about_stop(t):  # stop when we hit the next clearly-new section or sign-off
    return t.strip() == "" or t.strip().startswith("#") or _starts("www.upmc.com")(t)

BLOCK_STRIP_RULES = [
    ("STRIP_ABOUT",    "'About the University of Pittsburgh School of Medicine' boilerplate",
     _starts("About the University of Pittsburgh"), _about_stop),
    ("STRIP_ADDL_RES", "'Additional Resources' section",
     _starts("Additional Resources"), lambda t: t.strip() == "" or _starts("Pitt co-authors", "This research was supported", "#")(t)),
]

# Hyperlink classification by URL substring.
LINK_KEEP = ["doi.org", "nature.com", ".pitt.edu", "ncanda.org"]
LINK_UNLINK = ["upmc.com/services", "upmc.com/conditions", "upmc.com/locations",
               "upmc.com/pages", "upmc.com/media/news"]
def classify_link(url):
    u = (url or "").lower()
    if any(k in u for k in LINK_KEEP):
        return "KEEP"
    if "upmc.com" in u:
        return "UNLINK"   # default for any other UPMC link
    return "KEEP"

# Dateline prefix at start of the lede.
DATELINE_RE = re.compile(r"^[A-Z][A-Za-z.\s]+,\s+[A-Z][a-z]+\.?\s+\d{1,2},\s+\d{4}\s*[–—-]\s*")

# Hyphenation + dated-language substitutions (whole-word, case-insensitive).
HYPHEN_FIXES = {
    "longheld": "long-held", "risktaking": "risk-taking", "longterm": "long-term",
    "shortterm": "short-term", "decisionmaking": "decision-making",
    "wellknown": "well-known", "reallife": "real-life",
}
# "published today" → "published <Month Day>" using the release date parsed from
# the dateline/embargo (validated: published page reads "published June 11").
PUBLISHED_TODAY_RE = re.compile(r"\bpublished today\b", re.I)
RELEASE_DATE_RE = re.compile(
    r"\b(January|February|March|April|May|June|July|August|September|October|"
    r"November|December)\s+(\d{1,2}),?\s+\d{4}")

# Pitt HS web style DROPS academic degree credentials after names in body copy.
# Validated against the published medschool.pitt.edu article:
#   "Ashley Parr, Ph.D., research..."        -> "Ashley Parr, research..."
#   "Daniel Petrie, Ph.D., Finnegan Calabro" -> "Daniel Petrie, Finnegan Calabro"
#   "Duncan Clark, M.D., Ph.D.;"             -> "Duncan Clark;"
# Periodful forms (the way releases arrive) plus unambiguous ≥3-char periodless
# forms. 2-letter periodless forms (MD, MA, MS, BA, BS, DO, RN, JD) are EXCLUDED
# so we never eat a state code like "Bethesda, MD".
_DEGREE_TOKENS = ["Ph.D.", "M.D.", "M.S.", "M.A.", "M.P.H.", "M.F.A.", "Pharm.D.",
                  "D.O.", "Sc.D.", "Ed.D.", "D.N.P.", "M.S.N.", "M.B.A.", "J.D.",
                  "B.S.", "B.A.", "R.N.", "Dr.P.H.", "PhD", "MPH", "MFA",
                  "PharmD", "ScD", "EdD", "DNP", "MSN", "MBA", "DrPH"]
_DEG_ALT = "|".join(re.escape(t) for t in sorted(_DEGREE_TOKENS, key=len, reverse=True))
DEGREE_DROP_RE = re.compile(
    r",\s*(?:(?:%s)(?![A-Za-z]))(?:\s*,\s*(?:(?:%s)(?![A-Za-z])))*" % (_DEG_ALT, _DEG_ALT))

# Pitt prefix/decade spellings (validated against the published page). Literal,
# case-insensitive.
SPELLING_FIXES = {
    "non-invasive": "noninvasive", "co-author": "coauthor", "co-authors": "coauthors",
    "co‑author": "coauthor", "co‑authors": "coauthors",
    "mid-twenties": "mid-20s", "mid-thirties": "mid-30s",
    "early twenties": "early 20s", "late twenties": "late 20s",
    "early thirties": "early 30s", "late thirties": "late 30s",
}

# Pitt style: no space around em/en dashes.  Whole-document regex pass.
DASH_SPACING_RE = re.compile(r"\s*([—–])\s*")
# Pitt style: collapse runs of 2+ spaces to one.
DOUBLE_SPACE_RE = re.compile(r"  +")
# Pitt style: "8 percent" -> "8%".
PERCENT_RE = re.compile(r"\b(\d+(?:\.\d+)?)\s*percent\b", re.I)

# Paragraphs that are judgment calls -> flag only, never auto-strip.
FLAG_ONLY = [
    ("FLAG_FUNDING",  "Funding/grant paragraph (usually kept)", _starts("This research was supported")),
    ("FLAG_COAUTHOR", "Co-authors paragraph (usually kept)",    _starts("Pitt co-authors")),
]

# --------------------------------------------------------------------------- #
# Helpers                                                                       #
# --------------------------------------------------------------------------- #

def para_text(p):
    return p.text

def iter_body_paragraphs(doc):
    return list(doc.paragraphs)

def slugify(s):
    s = re.sub(r"[^A-Za-z0-9]+", "-", s.strip().lower()).strip("-")
    return re.sub(r"-{2,}", "-", s)[:60] or "release"

def delete_paragraph(p):
    el = p._element
    el.getparent().remove(el)
    p._p = p._element = None

def unlink_hyperlink(hyperlink_el):
    """Replace a <w:hyperlink> with its child runs (keep text, drop the link)."""
    parent = hyperlink_el.getparent()
    idx = list(parent).index(hyperlink_el)
    for child in list(hyperlink_el):
        if child.tag == qn("w:r"):
            # strip hyperlink styling so it reads as body text
            rpr = child.find(qn("w:rPr"))
            if rpr is not None:
                for tag in ("w:rStyle", "w:color", "w:u"):
                    e = rpr.find(qn(tag))
                    if e is not None:
                        rpr.remove(e)
            parent.insert(idx, child)
            idx += 1
    parent.remove(hyperlink_el)

# --------------------------------------------------------------------------- #
# DETECT                                                                        #
# --------------------------------------------------------------------------- #

def detect(path):
    doc = Document(path)
    paras = iter_body_paragraphs(doc)
    changes = []           # each: {id, category, label, type, detail, para_index}
    consumed = set()

    # Release date (for "published today" → "published <Month Day>").
    _alltext = "\n".join(para_text(p) for p in paras)
    _rd = RELEASE_DATE_RE.search(_alltext)
    release_date = f"{_rd.group(1)} {int(_rd.group(2))}" if _rd else None

    # Block strips first (they own a range of paragraphs).
    for cat, label, start_m, stop_m in BLOCK_STRIP_RULES:
        i = 0
        while i < len(paras):
            t = para_text(paras[i])
            if i not in consumed and t.strip() and start_m(t):
                j = i + 1
                block = [i]
                while j < len(paras):
                    tj = para_text(paras[j])
                    if stop_m(tj):
                        break
                    block.append(j); j += 1
                for k in block:
                    consumed.add(k)
                changes.append({"id": f"{cat}#{i}", "category": cat, "label": label,
                                "type": "strip_block",
                                "detail": para_text(paras[i])[:90],
                                "para_indices": block})
                i = j
            else:
                i += 1

    # Single-paragraph strips.
    for i, p in enumerate(paras):
        if i in consumed:
            continue
        t = para_text(p)
        if not t.strip():
            continue
        for cat, label, matcher in PARA_STRIP_RULES:
            try:
                hit = matcher(t)
            except Exception:
                hit = False
            if hit:
                consumed.add(i)
                changes.append({"id": f"{cat}#{i}", "category": cat, "label": label,
                                "type": "strip_para", "detail": t[:90],
                                "para_indices": [i]})
                break

    # Dateline.
    for i, p in enumerate(paras):
        if i in consumed:
            continue
        m = DATELINE_RE.match(para_text(p))
        if m:
            changes.append({"id": f"DATELINE#{i}", "category": "DATELINE",
                            "label": "Remove dateline prefix", "type": "dateline",
                            "detail": m.group(0), "para_indices": [i]})
            break  # only the lede

    # Hyperlinks.
    for i, p in enumerate(paras):
        if i in consumed:
            continue
        for h in p._element.findall(qn("w:hyperlink")):
            rid = h.get(qn("r:id"))
            url = doc.part.rels[rid].target_ref if rid in doc.part.rels else ""
            disp = "".join(node.text or "" for node in h.iter(qn("w:t")))
            action = classify_link(url)
            if action == "UNLINK":
                changes.append({"id": f"UNLINK#{i}#{rid}", "category": "UNLINK_UPMC",
                                "label": "Unlink UPMC marketing link (keep text)",
                                "type": "unlink", "detail": f"“{disp}” → {url}",
                                "para_indices": [i], "rid": rid})

    # Style: hyphenation + phrase fixes.
    for i, p in enumerate(paras):
        if i in consumed:
            continue
        t = para_text(p)
        low = t.lower()
        for bad, good in HYPHEN_FIXES.items():
            for m in re.finditer(rf"\b{re.escape(bad)}\b", t, flags=re.I):
                changes.append({"id": f"HYPHEN#{i}#{bad}#{m.start()}", "category": "STYLE_HYPHEN",
                                "label": "Restore hyphen", "type": "subword",
                                "detail": f"{m.group(0)} → {good}",
                                "para_indices": [i], "find": m.group(0), "repl": good})
        m = PUBLISHED_TODAY_RE.search(t)
        if m:
            repl = f"published {release_date}" if release_date else "published"
            changes.append({"id": f"DATED#{i}", "category": "STYLE_DATED",
                            "label": "Dated language → specific date", "type": "subphrase",
                            "detail": f"'{m.group(0)}' → '{repl}'",
                            "para_indices": [i], "find": m.group(0), "repl": repl})

    # Whole-document style passes (degrees, dash spacing, double spaces, percent).
    # Reported as one summary entry per category, applied as per-run regex/str
    # passes over the paragraphs that survive (skipping anything being stripped).
    live_text = " ".join(para_text(p) for i, p in enumerate(paras) if i not in consumed)

    deg_matches = DEGREE_DROP_RE.findall(live_text)
    if deg_matches:
        sample = ", ".join(s.strip().lstrip(",").strip() for s in deg_matches[:6])
        changes.append({"id": "STYLE_DEGREE", "category": "STYLE_DEGREE",
                        "label": "Drop academic degree credentials after names (HS web style)",
                        "type": "global_style", "pass": "degree",
                        "detail": f"{len(deg_matches)} credential group(s): {sample}",
                        "para_indices": []})

    spell_hits = []
    low_all = live_text.lower()
    for bad, good in SPELLING_FIXES.items():
        n = low_all.count(bad.lower())
        if n:
            spell_hits.append(f"{bad}→{good} ({n})")
    if spell_hits:
        changes.append({"id": "STYLE_SPELLING", "category": "STYLE_SPELLING",
                        "label": "Pitt prefix/decade spellings", "type": "global_style",
                        "pass": "spelling", "detail": ", ".join(spell_hits), "para_indices": []})

    n_dash = len([m for m in DASH_SPACING_RE.finditer(live_text)
                  if m.group(0) != m.group(1)])  # only spaced dashes
    if n_dash:
        changes.append({"id": "STYLE_DASH", "category": "STYLE_DASH",
                        "label": "Dashes: remove surrounding spaces (Pitt style)",
                        "type": "global_style", "pass": "dash",
                        "detail": f"{n_dash} spaced em/en dash(es)", "para_indices": []})

    n_sp = len(DOUBLE_SPACE_RE.findall(live_text))
    if n_sp:
        changes.append({"id": "STYLE_SPACING", "category": "STYLE_SPACING",
                        "label": "Collapse double spaces (Pitt style)",
                        "type": "global_style", "pass": "spacing",
                        "detail": f"{n_sp} run(s) of 2+ spaces", "para_indices": []})

    pct = PERCENT_RE.findall(live_text)
    if pct:
        changes.append({"id": "STYLE_PERCENT", "category": "STYLE_PERCENT",
                        "label": "Percent: 'N percent' → 'N%' (Pitt style)",
                        "type": "global_style", "pass": "percent",
                        "detail": ", ".join(f"{n} percent→{n}%" for n in pct[:6]),
                        "para_indices": []})

    # Flag-only judgment calls.
    for i, p in enumerate(paras):
        if i in consumed:
            continue
        t = para_text(p)
        for cat, label, matcher in FLAG_ONLY:
            if t.strip() and matcher(t):
                changes.append({"id": f"{cat}#{i}", "category": cat, "label": label,
                                "type": "flag", "detail": t[:90], "para_indices": [i]})

    plan = {"input": os.path.abspath(path),
            "headline": next((para_text(p) for p in paras if p.style and
                              p.style.name.lower().startswith("title")), None),
            "changes": changes}
    return plan

def print_plan(plan):
    by_cat = {}
    for c in plan["changes"]:
        by_cat.setdefault(c["category"], []).append(c)
    order = ["STRIP_CONTACT","STRIP_EMBARGO","STRIP_SUMMARY","STRIP_HASHES",
             "STRIP_MEDIA_FOOTER","STRIP_ABOUT","STRIP_ADDL_RES","DATELINE",
             "UNLINK_UPMC","STYLE_DEGREE","STYLE_SPELLING","STYLE_DASH","STYLE_SPACING",
             "STYLE_PERCENT","STYLE_HYPHEN","STYLE_DATED","FLAG_FUNDING","FLAG_COAUTHOR"]
    cats = [c for c in order if c in by_cat] + [c for c in by_cat if c not in order]
    print("\n=== PROPOSED CHANGES (nothing applied yet) ===\n")
    for cat in cats:
        items = by_cat[cat]
        flag = " [FLAG — review, not auto-applied]" if cat.startswith("FLAG") else ""
        print(f"[{cat}] {items[0]['label']}{flag}  ({len(items)})")
        for it in items:
            print(f"    - {it['detail']}")
        print()
    auto = [c for c in cats if not c.startswith("FLAG")]
    print("Approve with:  --approve " + (",".join(auto) if auto else "(none)"))
    print("Or:            --approve ALL   (applies every non-FLAG category)\n")

# --------------------------------------------------------------------------- #
# APPLY                                                                         #
# --------------------------------------------------------------------------- #

def apply(path, plan, approved, slug=None, outdir=None):
    doc = Document(path)
    paras = iter_body_paragraphs(doc)
    approved = set(approved)
    if "ALL" in approved:
        approved = {c["category"] for c in plan["changes"] if not c["category"].startswith("FLAG")}

    to_delete = set()
    applied = []
    global_passes = []   # ("degree"|"dash"|"spacing"|"percent")

    for c in plan["changes"]:
        cat = c["category"]
        if cat not in approved:
            continue
        if c["type"] == "global_style":
            global_passes.append(c["pass"])
            applied.append(c)
            continue
        if c["type"] in ("strip_para", "strip_block"):
            for idx in c["para_indices"]:
                to_delete.add(idx)
            applied.append(c)
        elif c["type"] == "dateline":
            p = paras[c["para_indices"][0]]
            if _strip_prefix_in_para(p, DATELINE_RE):
                applied.append(c)
        elif c["type"] in ("subword", "subphrase"):
            p = paras[c["para_indices"][0]]
            _replace_in_para(p, c["find"], c["repl"])
            applied.append(c)
        elif c["type"] == "unlink":
            p = paras[c["para_indices"][0]]
            for h in p._element.findall(qn("w:hyperlink")):
                if h.get(qn("r:id")) == c.get("rid"):
                    unlink_hyperlink(h)
            applied.append(c)

    # whole-document style passes over surviving paragraphs.
    def _style_fn(txt):
        if not txt:
            return txt
        if "degree" in global_passes:
            txt = DEGREE_DROP_RE.sub("", txt)
        if "spelling" in global_passes:
            for bad, good in SPELLING_FIXES.items():
                txt = re.sub(re.escape(bad), good, txt, flags=re.I)
        if "dash" in global_passes:
            txt = DASH_SPACING_RE.sub(lambda m: m.group(1), txt)
        if "spacing" in global_passes:
            txt = DOUBLE_SPACE_RE.sub(" ", txt)
        if "percent" in global_passes:
            txt = PERCENT_RE.sub(lambda m: m.group(1) + "%", txt)
        return txt

    if global_passes:
        for i, p in enumerate(paras):
            if i not in to_delete:
                _transform_para_text(p, _style_fn)

    # delete paragraphs last (indices already captured)
    for idx in sorted(to_delete, reverse=True):
        delete_paragraph(paras[idx])

    if slug is None:
        base = plan.get("headline") or os.path.splitext(os.path.basename(path))[0]
        slug = slugify(base)
    date = datetime.date.today().isoformat()
    outdir = outdir or os.path.dirname(os.path.abspath(path))
    outname = f"{date}-{slug}-web.docx"
    outpath = os.path.join(outdir, outname)
    doc.save(outpath)
    return outpath, applied

def _direct_runs(p):
    """Runs that are DIRECT children of the paragraph (excludes runs nested
    inside <w:hyperlink>), in document order. Editing only these means we never
    disturb hyperlink text or reorder anything."""
    from docx.text.run import Run
    return [Run(child, p) for child in p._element if child.tag == qn("w:r")]

def _para_has_hyperlink(p):
    return p._element.find(qn("w:hyperlink")) is not None

def _transform_para_text(p, fn):
    """Apply text-transform fn to a paragraph's own text.

    For a paragraph with NO hyperlinks we operate on the full concatenated text
    (so multi-run words like a split "co-authors" are caught), writing the result
    into the first run. For a paragraph WITH hyperlinks we transform each direct
    run independently — never moving text across runs — so link order/text stay
    intact (at the cost of missing a fix that straddles a run boundary)."""
    runs = _direct_runs(p)
    if not runs:
        return
    if _para_has_hyperlink(p):
        for r in runs:
            nt = fn(r.text)
            if nt != r.text:
                r.text = nt
    else:
        full = "".join(r.text for r in runs)
        nt = fn(full)
        if nt != full:
            runs[0].text = nt
            for r in runs[1:]:
                r.text = ""

def _replace_in_para(p, find, repl):
    """Case-insensitive replace within the paragraph's own runs, in place.

    Tries a single-run match first (the common case). If the phrase straddles
    consecutive direct runs, it edits that contiguous span only. Hyperlink runs
    are never touched, and no text is moved between runs."""
    runs = _direct_runs(p)
    fl = find.lower()

    # 1) single-run hit
    for r in runs:
        idx = r.text.lower().find(fl)
        if idx != -1:
            r.text = r.text[:idx] + repl + r.text[idx + len(find):]
            return True

    # 2) span across consecutive direct runs
    for start in range(len(runs)):
        acc = ""
        for end in range(start, len(runs)):
            acc += runs[end].text
            pos = acc.lower().find(fl)
            if pos != -1:
                # rebuild only this contiguous span; preserve leading/trailing text
                lead = acc[:pos]
                tail = acc[pos + len(find):]
                runs[start].text = lead + repl + tail
                for r in runs[start + 1:end + 1]:
                    r.text = ""
                return True
            if len(acc) > len(fl) + 200:
                break
    return False

def _strip_prefix_in_para(p, prefix_regex):
    """Remove a leading prefix (e.g., the dateline) from the start of the
    paragraph, editing only the leading direct runs."""
    runs = _direct_runs(p)
    acc = ""
    for i, r in enumerate(runs):
        acc += r.text
        m = prefix_regex.match(acc)
        if m and m.end() <= len(acc):
            remove = m.end()
            # walk runs again, deleting `remove` chars from the front
            for rr in runs:
                if remove <= 0:
                    break
                if len(rr.text) <= remove:
                    remove -= len(rr.text)
                    rr.text = ""
                else:
                    rr.text = rr.text[remove:]
                    remove = 0
            return True
        # only the very start matters; stop once the run text is non-prefix-y
        if len(acc) > 120:
            break
    return False

# --------------------------------------------------------------------------- #
# propose: capture a writer correction as a portable rule proposal             #
# --------------------------------------------------------------------------- #

# The categories the change list already knows about. A proposal may name a new
# one (writers find gaps the rules don't cover yet) — that's allowed, we just
# flag it so the maintainer notices a genuinely new rule type.
KNOWN_CATEGORIES = {
    "STRIP_CONTACT", "STRIP_EMBARGO", "STRIP_SUMMARY", "STRIP_HASHES",
    "STRIP_MEDIA_FOOTER", "STRIP_ABOUT", "STRIP_ADDL_RES", "DATELINE",
    "UNLINK_UPMC", "STYLE_DEGREE", "STYLE_SPELLING", "STYLE_DASH",
    "STYLE_SPACING", "STYLE_PERCENT", "STYLE_DATED", "STYLE_HYPHEN",
    "FLAG_FUNDING", "FLAG_COAUTHOR",
}

def propose(category, before, after, example=None, submitter=None,
            source=None, note=None, out=None):
    """Append a writer correction to the team's rule-proposal files.

    This deliberately changes NO rule. It records a suggestion for the maintainer
    to review and merge, so every teammate's correction lands in one consistent
    format regardless of who ran it. Two files are written side by side next to
    the release (or wherever --out points): a human-readable `.md` the maintainer
    skims, and a `.jsonl` mirror that's easy to merge or script against later.
    """
    category = category.strip().upper()
    known = category in KNOWN_CATEGORIES
    date = datetime.date.today().isoformat()

    md_path = out or os.path.join(os.getcwd(), "upmc-rule-proposals.md")
    jsonl_path = os.path.splitext(md_path)[0] + ".jsonl"

    record = {
        "date": date,
        "category": category,
        "known_category": known,
        "before": before,
        "after": after,
        "example": example or "",
        "submitter": submitter or "",
        "source": os.path.basename(source) if source else "",
        "note": note or "",
        "status": "pending",
    }

    # Human-readable markdown block (append; write a header the first time).
    header = ""
    if not os.path.exists(md_path):
        header = (
            "# UPMC-release-to-web — rule proposals\n\n"
            "Writer corrections captured for the maintainer to review and merge. "
            "Nothing here changes a rule until the maintainer folds it into the "
            "shared rule set (see the skill's \"For the maintainer\" section). "
            "Mark or delete a block once it has been handled.\n"
        )
    tag = "" if known else "  ⚠ new/unknown category — maintainer, confirm this rule type"
    block = [f"## {date} — {category}{tag}",
             f"- **Before:** `{before}`",
             f"- **After:** `{after}`"]
    if example:
        block.append(f"- **Example:** {example}")
    if record["source"]:
        block.append(f"- **Source release:** {record['source']}")
    if submitter:
        block.append(f"- **Submitted by:** {submitter}")
    if note:
        block.append(f"- **Why:** {note}")
    block.append("- **Status:** pending")
    with open(md_path, "a") as f:
        if header:
            f.write(header)
        f.write("\n" + "\n".join(block) + "\n")

    # Machine-readable JSONL mirror.
    with open(jsonl_path, "a") as f:
        f.write(json.dumps(record) + "\n")

    return md_path, jsonl_path, known

# --------------------------------------------------------------------------- #
# CLI                                                                           #
# --------------------------------------------------------------------------- #

def main():
    ap = argparse.ArgumentParser()
    sub = ap.add_subparsers(dest="cmd", required=True)

    d = sub.add_parser("detect")
    d.add_argument("input")
    d.add_argument("--out", default=None)

    a = sub.add_parser("apply")
    a.add_argument("input")
    a.add_argument("--plan", required=True)
    a.add_argument("--approve", required=True, help="comma list of categories or ALL")
    a.add_argument("--slug", default=None)
    a.add_argument("--outdir", default=None)

    pr = sub.add_parser("propose", help="record a writer correction for the maintainer to merge")
    pr.add_argument("--category", required=True, help="e.g. STYLE_SPELLING, UNLINK_UPMC, or a new UPPER_SNAKE name")
    pr.add_argument("--before", required=True, help="the text/behavior as the skill produced it")
    pr.add_argument("--after", required=True, help="what it should have been")
    pr.add_argument("--example", default=None, help="a short in-context before -> after example")
    pr.add_argument("--submitter", default=None, help="who caught it")
    pr.add_argument("--source", default=None, help="the release .docx the correction came from")
    pr.add_argument("--note", default=None, help="why this is the right call")
    pr.add_argument("--out", default=None, help="proposals .md path (default: ./upmc-rule-proposals.md)")

    args = ap.parse_args()

    if args.cmd == "detect":
        plan = detect(args.input)
        outp = args.out or (os.path.splitext(args.input)[0] + ".plan.json")
        with open(outp, "w") as f:
            json.dump(plan, f, indent=2)
        print_plan(plan)
        print(f"Plan written to: {outp}")
    elif args.cmd == "apply":
        with open(args.plan) as f:
            plan = json.load(f)
        approved = [c.strip() for c in args.approve.split(",") if c.strip()]
        outpath, applied = apply(args.input, plan, approved, args.slug, args.outdir)
        print(f"Applied {len(applied)} change(s). Wrote copy:\n  {outpath}")
        print(f"Original untouched: {os.path.abspath(args.input)}")
    elif args.cmd == "propose":
        md_path, jsonl_path, known = propose(
            args.category, args.before, args.after, args.example,
            args.submitter, args.source, args.note, args.out)
        note = "" if known else f"  (heads up: '{args.category.strip().upper()}' isn't a known category — fine if that's intended)"
        print(f"Proposal recorded.{note}")
        print(f"  {md_path}")
        print(f"  {jsonl_path}")
        print("Send these to the skill maintainer to fold into the next version.")

if __name__ == "__main__":
    main()
