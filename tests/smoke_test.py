#!/usr/bin/env python3
"""CI smoke test for the pitt-hs marketplace.

Runs on every push/PR. Fails the build if the marketplace/plugin manifests are
malformed or if the upmc-release-to-web transform engine stops behaving, so a
broken release never reaches the team's `/plugin update`.

Checks:
  1. marketplace.json + plugin.json are valid JSON with the expected fields, and
     the marketplace's plugin `source` path actually resolves to a plugin.
  2. plugin.json has a version (updates are triggered by it — a missing/blank
     version would strand the team on the old copy).
  3. transform.py runs the full detect -> apply -> propose cycle: the original is
     never modified, the dated copy IS transformed, and a proposal is captured.

Only dependency: python-docx (matches the skill's own requirement).
"""
import json, os, re, subprocess, sys, tempfile

REPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
MARKET = os.path.join(REPO, ".claude-plugin", "marketplace.json")

failures = []
def check(cond, msg):
    print(("  ok  " if cond else "FAIL  ") + msg)
    if not cond:
        failures.append(msg)

print("== manifests ==")
mp = json.load(open(MARKET))
check(isinstance(mp.get("plugins"), list) and mp["plugins"], "marketplace.json lists at least one plugin")
plugin_dir = None
for entry in mp["plugins"]:
    src = entry.get("source", "")
    check(src.startswith("./"), f"plugin source is a repo-relative path: {src!r}")
    resolved = os.path.normpath(os.path.join(REPO, src))
    pj = os.path.join(resolved, ".claude-plugin", "plugin.json")
    check(os.path.isfile(pj), f"plugin.json exists at {os.path.relpath(pj, REPO)}")
    if os.path.isfile(pj):
        d = json.load(open(pj))
        check(bool(re.fullmatch(r"[a-z0-9-]+", d.get("name", ""))), f"plugin name is kebab-case: {d.get('name')!r}")
        check(bool(re.fullmatch(r"\d+\.\d+\.\d+", d.get("version", ""))), f"plugin version is semver: {d.get('version')!r}")
        if d.get("name") == "pitt-hs-content-ops":
            plugin_dir = resolved

print("== transform.py end-to-end ==")
transform = os.path.join(plugin_dir or "", "skills", "upmc-release-to-web", "scripts", "transform.py")
check(os.path.isfile(transform), "transform.py found in the plugin")

if os.path.isfile(transform):
    try:
        from docx import Document
    except ImportError:
        print("FAIL  python-docx not installed"); sys.exit(1)

    with tempfile.TemporaryDirectory() as tmp:
        src = os.path.join(tmp, "release.docx")
        doc = Document()
        for line in [
            "EMBARGOED FOR RELEASE UNTIL 11 A.M. ET, JUNE 11, 2026",
            "Contact: Jane Doe",
            "PITTSBURGH, June 11, 2026 – Lower dopamine may drive teen risk-taking, "
            "published today by University of Pittsburgh researchers.",
            "The effect was 8 percent stronger among the non-invasive imaging group.",
            "# # #",
            "www.upmc.com/media",
        ]:
            doc.add_paragraph(line)
        doc.save(src)

        def run(*args):
            return subprocess.run([sys.executable, transform, *args],
                                  cwd=tmp, capture_output=True, text=True)

        r = run("detect", src)
        check(r.returncode == 0 and "PROPOSED CHANGES" in r.stdout, "detect runs and lists changes")
        plan = src.replace(".docx", ".plan.json")
        check(os.path.isfile(plan), "detect writes a plan file")

        r = run("apply", src, "--plan", plan, "--approve", "ALL", "--slug", "smoke")
        check(r.returncode == 0, "apply runs")
        copies = [f for f in os.listdir(tmp) if f.endswith("-smoke-web.docx")]
        check(len(copies) == 1, "apply writes exactly one dated -web copy")
        orig_text = "\n".join(p.text for p in Document(src).paragraphs)
        check("EMBARGOED" in orig_text, "original is left untouched (still has embargo line)")
        if copies:
            copy_text = "\n".join(p.text for p in Document(os.path.join(tmp, copies[0])).paragraphs)
            check("EMBARGOED" not in copy_text, "copy has the embargo line stripped")
            check("noninvasive" in copy_text and "8%" in copy_text, "copy has style fixes applied")

        r = run("propose", "--category", "STYLE_SPELLING", "--before", "healthcare",
                "--after", "health care", "--note", "smoke test",
                "--out", os.path.join(tmp, "props.md"))
        check(r.returncode == 0 and os.path.isfile(os.path.join(tmp, "props.md")), "propose captures a proposal (.md)")
        check(os.path.isfile(os.path.join(tmp, "props.jsonl")), "propose writes the .jsonl mirror")

print()
if failures:
    print(f"SMOKE TEST FAILED: {len(failures)} problem(s)")
    sys.exit(1)
print("SMOKE TEST PASSED")
